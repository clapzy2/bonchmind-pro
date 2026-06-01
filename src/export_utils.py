"""
export_utils.py - экспорт ответов BonchMind Pro в DOCX.
"""

import re

from datetime import datetime
from pathlib import Path

from docx import Document


EXPORT_DIR = Path("exports")

def make_safe_filename(text):
    """Сделать строку безопасной для имени файла."""
    text = str(text).strip().lower()
    text = re.sub(r"[^\wа-яёА-ЯЁ]+", "_", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("_")[:80] or "export"

def _clean_markdown(text):
    """Убирает простую markdown-разметку для DOCX."""
    text = str(text)
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = text.replace("**", "")
    text = text.replace("---", "")
    return text.strip()

def _message_to_text(content):
    """Преобразовать сообщение Gradio в обычный текст."""
    if isinstance(content, str):
        return content.strip()

    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                parts.append(str(item.get("text", "")).strip())
            else:
                parts.append(str(item).strip())
        return "\n".join(part for part in parts if part)

    if isinstance(content, dict):
        return str(content.get("text", "")).strip()

    return str(content).strip()


def _extract_last_qa(history):
    """Получить последний вопрос и ответ из истории Gradio."""
    if not history:
        return "", ""

    last_question = ""
    last_answer = ""

    for msg in reversed(history):
        role = msg.get("role", "")
        content = _message_to_text(msg.get("content", ""))

        if role == "assistant" and not last_answer:
            last_answer = content
        elif role == "user" and not last_question:
            last_question = content
            break

    return last_question, last_answer

def export_text_to_docx(title, content, prefix="bonchmind_export", name_parts=None):
    """Экспортирует произвольный текст в DOCX."""
    if not content:
        return None

    EXPORT_DIR.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    name_suffix = ""
    if name_parts:
        safe_parts = [make_safe_filename(part) for part in name_parts if part]
        if safe_parts:
            name_suffix = "_" + "_".join(safe_parts)

    filename = EXPORT_DIR / f"{prefix}{name_suffix}_{timestamp}.docx"

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    for raw_block in str(content).split("\n"):
        raw_block = raw_block.strip()
        if not raw_block:
            continue

        # Markdown-заголовки превращаем в заголовки Word
        if raw_block.startswith("### "):
            doc.add_heading(_clean_markdown(raw_block), level=3)
        elif raw_block.startswith("## "):
            doc.add_heading(_clean_markdown(raw_block), level=2)
        elif raw_block.startswith("# "):
            doc.add_heading(_clean_markdown(raw_block), level=1)
        else:
            block = _clean_markdown(raw_block)
            if block:
                doc.add_paragraph(block)

    doc.save(filename)
    return str(filename)

def export_last_answer_to_docx(history):
    """
    Экспортирует последний вопрос и ответ в DOCX.
    Возвращает путь к созданному файлу.
    """
    question, answer = _extract_last_qa(history)

    if not question or not answer:
        return None

    EXPORT_DIR.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = EXPORT_DIR / f"bonchmind_answer_{timestamp}.docx"

    doc = Document()

    doc.add_heading("BonchMind Pro — экспорт ответа", level=1)
    doc.add_paragraph(f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    doc.add_heading("Вопрос", level=2)
    doc.add_paragraph(question)

    doc.add_heading("Ответ", level=2)
    for raw_block in answer.split("\n"):
        raw_block = raw_block.strip()
        if not raw_block:
            continue

        if raw_block.startswith("### "):
            doc.add_heading(_clean_markdown(raw_block), level=3)
        elif raw_block.startswith("## "):
            doc.add_heading(_clean_markdown(raw_block), level=2)
        elif raw_block.startswith("# "):
            doc.add_heading(_clean_markdown(raw_block), level=1)
        else:
            block = _clean_markdown(raw_block)
            if block:
                doc.add_paragraph(block)

    doc.save(filename)
    return str(filename)